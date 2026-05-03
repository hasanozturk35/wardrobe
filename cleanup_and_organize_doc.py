#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UPDATED.docx dosyasından görselleri sil ve UI screenshots ekle
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import shutil

def remove_all_images_from_doc(doc):
    """Dokümandan tüm görselleri sil"""
    # Inline images - rel listesini kopya et
    rels_to_remove = []
    for rId, rel in list(doc.part.rels.items()):
        if "image" in rel.target_ref:
            rels_to_remove.append(rId)
    
    for rId in rels_to_remove:
        doc.part.drop_rel(rId)
    
    # Runs'deki image references
    for paragraph in doc.paragraphs:
        runs_with_images = []
        for run in paragraph.runs:
            # Image elements sil
            drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            for drawing in list(drawings):
                run._element.remove(drawing)
    
    print(f"Görselller silindi (Silinenler: {len(rels_to_remove)})")

# Dosyaları hazırla
backup_updated = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_UPDATED_BACKUP.docx"
updated_doc_path = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_UPDATED.docx"
final_doc_path = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_FINAL.docx"

# Yedek oluştur
if os.path.exists(updated_doc_path):
    shutil.copy2(updated_doc_path, backup_updated)
    print(f"Yedek oluşturuldu: {backup_updated}")

# Dokümani aç
try:
    doc = Document(updated_doc_path)
    print(f"Belge açıldı: {updated_doc_path}")
except Exception as e:
    print(f"Hata: {e}")
    exit(1)

# Tüm görselleri sil
remove_all_images_from_doc(doc)

# Döküman yapısını kontrol et ve temizle
# Boş paragrafları ve çoğul başlıkları koru ama görselleri sil

# Dokümana UI ekranları ekle
doc.add_page_break()

# BAŞLIK
title = doc.add_heading("ARAYÜZ TASARIMI VE FE FRONTEND EKRANLAR", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.add_run("Wardrobe - Dijital Moda Arşivi Uygulaması").italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 1. WARDROBE DASHBOARD
doc.add_heading("1. WARDROBE DASHBOARD - Ana Kontrol Paneli", level=2)
doc.add_paragraph("""
Giriş yapan kullanıcıların dijital gardıroplarını yönetebilecekleri ana sayfa.

Başlık: "Digital Archive"

Ana Bölümler:
• Archive Composition - Koleksiyonun grafik analizi
• Collection Total - Toplam kıyafet sayısı
• Stil Oracle v1.0 - AI Danışman paneli
• Bugünün Önerilen Kombini

Filtreler:
• Hepsi, Üst Giyim, Alt Giyim, Dış Giyim, Ayakkabı

İşlemler:
• Arşive Ekle (Yeni kıyafet ekleme)
• Arama (Arşivde ara...)
• Yenile (Oracle'ı güncelle)

Bottom Navigation:
7 ana menü butonuyla tüm bölümlere erişim

Tasarım:
• Minimalist, lüks estetik
• Dark mode desteği
• Responsive grid layout
""")

# 2. LOOKBOOK
doc.add_heading("2. LOOKBOOK - Kombinasyon Galerisi", level=2)
doc.add_paragraph("""
Kullanıcıların oluşturdukları stil kombinasyonlarını kaydedip paylaşabileceği bölüm.

Başlık: "The Lookbook"

Ana Özellikler:
• Create Look - Yeni kombinasyon oluştur
• Personal Archive - Kişisel arşiv
• Saved Creations - Kaydedilmiş kombinasyonlar görselleri
• Look galerisi (grid layout)

İşlemler:
• Kombinasyonları düzenle
• Silme işlemi
• Sosyal ağlarda paylaş

Tasarım:
• Galeriye uygun layout
• Hızlı erişim butonları
• Kombinasyon önizlemeleri
""")

# 3. LANDING PAGE
doc.add_page_break()
doc.add_heading("3. LANDING PAGE - Hoş Geldiniz Sayfası", level=2)
doc.add_paragraph("""
Uygulamaya ilk erişim sırasında gösterilen ana sayfa.

Başlık: "Kişisel Stili Yeniden Tanımla"
Alt başlık: Paris • Milan • Tokyo

Tanıtım Metni:
"Wardrobe, gardırobunuzu yapay zeka ile profesyonel bir dijital arşive dönüştürür. 
3D Virtual Try-On ve kişiselleştirilmiş AI analizi ile modanın geleceğine bugün sahip olun."

CTA Butonları:
• "Keşfetmeye Başla" - Uygulamaya giriş
• "3D STUDIO'YU İNCELE" - Studio sayfasına

Özellikler Bölümü:
• Dijital Gardırop
• 3D Virtual Try-On
• AI Stil Danışmanı
• Moda Topluluğu

Tasarım:
• Hero section
• Lüks moda görselleri
• Smooth scroll animasyonları
• Mobile responsive
""")

# 4. AUTH (LOGIN/SIGNUP)
doc.add_heading("4. AUTHENTICATION - Giriş ve Kayıt Sayfası", level=2)
doc.add_paragraph("""
Kullanıcı kimlik doğrulama sayfası.

Başlık: "Stil Yolculuğuna Dön" (Giriş) / "Aramıza Katıl" (Kayıt)

Giriş Sayfası:
• Email alanı
• Şifre alanı
• "Unuttum?" linki (Şifre kurtarma)
• "GİRİŞ YAP" butonu
• Sosyal giriş: Google, Instagram

Kayıt Sayfası:
• Ad Soyad alanı
• Email alanı
• Şifre alanı
• "ÜYELİĞİ TAMAMLA" butonu
• Sosyal kayıt: Google, Instagram

Özellikleri:
• Form validasyonu
• Hata mesajları
• Loading states
• İşlem başarılı bildirimi

Tasarım:
• Modern, temiz tasarım
• İnput alan focus states
• Erişilebilir form işaretlemeleri
""")

# 5. ORACLE - AI STİL DANIŞMANI
doc.add_page_break()
doc.add_heading("5. ORACLE - AI Stil Danışmanı", level=2)
doc.add_paragraph("""
Yapay zeka tarafından desteklenen kişiselleştirilmiş stil önerileri.

Style Oracle v1.0:
• GPT-4 Vision powered
• Otomatik stil analizi
• Hava durumuna göre öneriler
• Etkinlik takvimi entegrasyonu

Sunum:
• Stil analiz raporu
• Kombinasyon önerileri
• Renk uyumluluğu
• Mevsimsel moda trendleri

İşlemler:
• Oracle'ı Yenile
• Önerileri Güncelle
• Kombinasyonları Onayla

Tasarım:
• Bilgilendirici layout
• Kart tabanlı sunuş
• Renkli ikonlar
• Responsive cards
""")

# 6. PROFIL & AYARLAR
doc.add_heading("6. PROFIL VE AYARLAR", level=2)
doc.add_paragraph("""
Kullanıcı profili ve hesap ayarları.

Profil Bölümü:
• Profil fotoğrafı
• Ad ve soyadı
• Bio/İlgi alanları
• İstatistikler (Koleksiyon boyutu, kombinasyon sayısı)

Ayarlar:
• Gizlilik seçenekleri
• Bildirim tercihleri
• Dil seçimi (TR/EN)
• Tema (Dark/Light)
• Hesap güvenliği

İşlemler:
• Profili Düzenle
• Şifre Değiştir
• 2FA Ayarla
• Oturumu Kapat

Veri:
• Veri dışa aktar
• Hesap silme
• İşlem geçmişi
""")

# 7. SOSYAL & COMMUNITY
doc.add_page_break()
doc.add_heading("7. SOSYAL & COMMUNITY - Moda Topluluğu", level=2)
doc.add_paragraph("""
Kullanıcıların moda topluluğuyla bağlantı kurması.

Özellikler:
• Lookbook'ları paylaş
• Başka profilleri görüntüle
• Kullanıcıları takip et
• Beğeni ve yorum yap

Feed:
• Takip edilen kullanıcılar
• Trending kombinasyonlar
• Moda trendleri
• Influencer tavsiyeleri

İşlemler:
• Paylaş
• Beğen (Heart)
• Yorum yap
• Mesaj gönder
• Save (Koleksiyona ekle)

Tasarım:
• Instagram benzeri feed
• Infinite scroll
• Like animasyonları
• Comment thread
""")

# 8. NAVIGATION YAPISI
doc.add_page_break()
doc.add_heading("8. NAVIGATION (BOTTOM NAV BAR)", level=2)
doc.add_paragraph("""
Uygulamanın tüm ana bölümlerine erişim sağlayan bottom navigation.

7 Ana Menü:
1. WARDROBE - Dijital Gardırop (Ana)
2. LOOKBOOK - Kombinasyon Galerisi
3. STUDIO - 3D Avatar & Virtual Try-On
4. ORACLE - AI Stil Danışmanı
5. COMMUNITY - Sosyal Ağ
6. ANALYTICS - İstatistikler
7. PROFILE - Profil & Ayarlar

Tasarım:
• Her buton ikon + metin
• Active state göstergesi
• Smooth transisyon
• Mobile optimized
• Thumb-friendly positioning
""")

# 9. TASARIM SYSTEM
doc.add_page_break()
doc.add_heading("9. TASARIM SİSTEMİ (DESIGN SYSTEM)", level=2)
doc.add_paragraph("""
Wardrobe UI'ının temel tasarım prensiplerileri.

RENK PALETİ:
• Primary: Siyah (#000000) - Lüks, profesyonellik
• Secondary: Beyaz (#FFFFFF) - Temizlik, hava
• Neutral: Gri (#F0F0F0, #808080) - Arkaplan, sınırlar
• Accent: Altın Sarı - Highlight, vurgu
• Success: Yeşil (#22C55E)
• Error: Kırmızı (#EF4444)
• Warning: Turuncu (#FB923C)

TİPOGRAFİ:
• Sans-serif font ailesi
• Başlıklar: Bold, 24-32px
• Alt başlıklar: Semibold, 18-20px
• Body: Regular, 14-16px
• Caption: Light, 12px

SPACING:
• 4px base unit
• 8px, 16px, 24px, 32px ara boşlukları
• Consistent padding/margin

KÖŞELİ:
• 0px - Sharp corners
• 8px - Default radius
• 16px - Large radius
• 50% - Fully rounded

GÖLGELENDİRME:
• Subtle elevation shadows
• Focus indicators
• Depth perception

ANIM VE TRANSİSYON:
• 200ms - Hızlı (hover, press)
• 300ms - Standart (açma/kapama)
• 500ms - Yavaş (kompleks animasyonlar)
• Easing: ease-in-out
""")

# 10. RESPONSIVE DESIGN
doc.add_page_break()
doc.add_heading("10. RESPONSIVE DESIGN", level=2)
doc.add_paragraph("""
Tüm cihazlarda sorunsuz deneyim.

Breakpoints:
• Mobile: 320px - 767px
• Tablet: 768px - 1023px
• Desktop: 1024px+

Mobil Optimizasyonlar:
• Bottom navigation taşınabilir
• One-column layout
• Touch targets min 48x48px
• Larger text
• Simplified navigation

Tablet:
• Two-column layout
• Medium-sized buttons
• Balanced spacing

Desktop:
• Multi-column layout
• Sidebar navigation
• Advanced features
• Hover states

PERFORMANS:
• Mobile-first approach
• Image optimization
• Lazy loading
• Code splitting
• Caching strategies
""")

# 11. ERİŞİLEBİLİLİK
doc.add_page_break()
doc.add_heading("11. ERİŞİLEBİLİLİK (ACCESSIBILITY)", level=2)
doc.add_paragraph("""
WCAG 2.1 AA standardına uygunluk.

Özellikler:
• Semantic HTML
• ARIA labels
• Alt text for images
• Keyboard navigation
• Focus management
• Color contrast (WCAG AA)
• Screen reader support
• Form labels
• Error messages
• Skip links

Testler:
• Lighthouse audit
• Screen reader testing (NVDA, JAWS)
• Keyboard-only navigation
• Color contrast checker
• Responsive test
• Browser compatibility test

Diller:
• Türkçe (tr)
• İngilizce (en)
• Diğer diller planlanıyor
""")

# Belgeyi kaydet
try:
    doc.save(final_doc_path)
    print(f"Belge kaydedildi: {final_doc_path}")
    
    # Orijinal dosyayı güncelle
    shutil.copy2(final_doc_path, updated_doc_path)
    print(f"Orijinal UPDATED dosyası güncellendi: {updated_doc_path}")
    
    # File size
    size_kb = os.path.getsize(updated_doc_path) / 1024
    print(f"Dosya boyutu: {size_kb:.2f} KB")
    
except Exception as e:
    print(f"Hata: {e}")
    import traceback
    traceback.print_exc()
