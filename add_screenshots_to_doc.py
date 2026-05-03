#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UPDATED.docx dosyasına arayüz screenshot'larını ekle
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io
from PIL import Image

# Dosya yolları
updated_doc = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_UPDATED.docx"

# Dokümani aç
doc = Document(updated_doc)

# Dokümanda "ARAYÜZ TASARIMI VE" başlıktan sonra resimleri ekle
insert_index = None
for i, para in enumerate(doc.paragraphs):
    if "ARAYÜZ TASARIMI" in para.text:
        insert_index = i
        break

# Screenshot'ların bulunacağı kısaltılmış açıklamalar
screenshots_info = [
    {
        "title": "1. WARDROBE DASHBOARD",
        "description": "Kullanıcının dijital gardıroplarını yönetebilecekleri ana sayfa. 'Digital Archive' başlığıyla, koleksiyon analiz grafiği, filtreler, ve AI stil önerileri yer almaktadır."
    },
    {
        "title": "2. LOOKBOOK",
        "description": "'The Lookbook' sayfası, kullanıcıların oluşturdukları stil kombinasyonlarını kaydedip paylaşabileceği bölüm. 'Create Look' butonu ile yeni kombinasyon oluşturulabilir."
    },
    {
        "title": "3. LANDING PAGE (ANA SAYFA)",
        "description": "Uygulamaya ilk erişim sırasında gösterilen hoş geldiniz sayfası. 'Kişisel Stili Yeniden Tanımla' başlığı ve özellikleri tanıtan bölümler bulunmaktadır."
    },
    {
        "title": "4. AUTHENTICATION (LOGIN/SIGNUP)",
        "description": "Kullanıcı giriş ve kayıt sayfası. Email/Şifre alanları, 'Unuttum?' linki, ve sosyal medya ile giriş (Google, Instagram) seçenekleri yer almaktadır."
    },
    {
        "title": "5. ORACLE (AI STİL DANIŞMANI)",
        "description": "GPT-4 Vision tarafından desteklenen AI stil danışmanı. Bottom navigation bar tüm sayfaları birleştiren ana menüyü göstermektedir."
    }
]

# Screenshot yolları
screenshots_paths = [
    (r"c:\Users\LENOVO\son bitirme\screenshots\1_wardrobe_dashboard.png", "Wardrobe Dashboard"),
    (r"c:\Users\LENOVO\son bitirme\screenshots\2_lookbook.png", "Lookbook"),
    (r"c:\Users\LENOVO\son bitirme\screenshots\3_landing_page.png", "Landing Page"),
    (r"c:\Users\LENOVO\son bitirme\screenshots\4_auth.png", "Auth Page"),
]

# Sayfayı ekle ve screenshot'ları koy
for info in screenshots_info:
    # Başlık
    heading = doc.add_heading(info["title"], level=3)
    
    # Açıklama
    doc.add_paragraph(info["description"])
    
    doc.add_paragraph()

doc.add_page_break()

# Screenshot'ları ekle
doc.add_heading("ARAYÜZ EKRAN GÖRÜNTÜLERİ (SCREENSHOTS)", level=2)
doc.add_paragraph("Wardrobe uygulamasının ayrıntılı ekran görüntüleri aşağıda gösterilmektedir.")
doc.add_paragraph()

screenshot_count = 0
for img_path, caption in screenshots_paths:
    if os.path.exists(img_path):
        try:
            # Resim boyutunu kontrol et
            with Image.open(img_path) as img:
                width, height = img.size
                # Aspect ratio'yu koru
                doc.add_picture(img_path, width=Inches(5.5))
                screenshot_count += 1
                
                # Başlık ekle
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Caption ekle
                caption_para = doc.add_paragraph(caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.runs[0].font.size = Pt(10)
                caption_para.runs[0].font.italic = True
                
                doc.add_paragraph()
                
        except Exception as e:
            print(f"Resim eklenirken hata: {img_path} - {e}")
    else:
        print(f"Resim bulunamadı: {img_path}")

# Belgeyi kaydet
try:
    doc.save(updated_doc)
    print(f"✅ {screenshot_count} arayüz ekran görüntüsü eklendi!")
    print(f"📄 Dosya kaydedildi: {updated_doc}")
    print(f"📊 Dosya boyutu: {os.path.getsize(updated_doc) / 1024:.2f} KB")
except Exception as e:
    print(f"❌ Hata: {e}")
    import traceback
    traceback.print_exc()
