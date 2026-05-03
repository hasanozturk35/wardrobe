#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UPDATED.docx dosyasına mevcut görselleri ekle
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil

# Dosya yolları
source_doc = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_UPDATED_BACKUP.docx"
updated_doc = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_WITH_IMAGES.docx"

# Eğer source yoksa, backup'dan kopyala
if not os.path.exists(source_doc):
    print(f"Source dosya yok, FINAL.docx'ı kullanıyorum...")
    source_doc = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_FINAL.docx"

if os.path.exists(source_doc):
    shutil.copy2(source_doc, updated_doc)
    print(f"Dosya kopyalandı: {updated_doc}")
else:
    print("Uygun kaynak dosya bulunamadı!")
    exit(1)

# Dokümani aç
doc = Document(updated_doc)

# Sayfayı ekle
doc.add_page_break()
doc.add_heading("ARAYÜZ EKRAN GÖRÜNTÜLERİ VE VİZÜAL ÖRNEKLER", level=2)
doc.add_paragraph("""
Wardrobe uygulamasının tüm arayüz ekranları React ve TailwindCSS ile tasarlanmıştır. 
Aşağıda projenin kullanılan stil ve görsel öğeleri gösterilmektedir.
""")

# Mevcut görselleri ekle
assets_path = r"c:\Users\LENOVO\son bitirme\frontend\src\assets"
uploads_path = r"c:\Users\LENOVO\son bitirme\backend\uploads"

added_count = 0

# Assets'teki görselleri ekle
if os.path.exists(assets_path):
    doc.add_heading("UI Tasarım Öğeleri", level=3)
    
    assets = [
        ("chic_fashion_shoot_photo_1772833711579.png", "Premium Moda Görseli - Landing Page"),
        ("fashion_avatar_render_1772833388818.png", "3D Avatar Render - Avatar Modülü"),
        ("quiet_luxury_wardrobe_bg_1772833697604.png", "Quiet Luxury Arka Planı"),
        ("premium_fashion_visual_1772833612094.png", "Premium Tasarım Görseli"),
        ("pinterest_fashion_couple_editorial_1772833854361.png", "Moda Editöryeli - Lookbook Referansı"),
    ]
    
    for filename, caption in assets:
        filepath = os.path.join(assets_path, filename)
        if os.path.exists(filepath):
            try:
                doc.add_picture(filepath, width=Inches(4.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Caption
                cap_para = doc.add_paragraph(caption)
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_para.runs[0].font.size = Pt(9)
                cap_para.runs[0].font.italic = True
                
                doc.add_paragraph()
                added_count += 1
            except Exception as e:
                print(f"Görüntü eklenirken hata: {filename} - {e}")

# Ürün Katalog Görselleri
doc.add_heading("Ürün Katalog Örnekleri", level=3)
doc.add_paragraph("Wardrobe'da yönetilen kıyafet ürünlerinin örnek görüntüleri:")

if os.path.exists(uploads_path):
    img_count = 0
    for filename in sorted(os.listdir(uploads_path)):
        if filename.startswith("opt_") and filename.endswith(".png") and img_count < 6:
            filepath = os.path.join(uploads_path, filename)
            try:
                doc.add_picture(filepath, width=Inches(2.0))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_count += 1
                added_count += 1
                
                # Her 3 resime bir satır ayır
                if img_count % 3 == 0:
                    doc.add_paragraph()
            except Exception as e:
                print(f"Ürün görüntüsü eklenirken hata: {filename}")

# Tasarım Sistemi Detayları
doc.add_page_break()
doc.add_heading("UI/UX TASARIM STANDARTLARI", level=2)

doc.add_heading("Renk Şeması (Color Scheme)", level=3)
doc.add_paragraph("""
Primary Colors:
• Siyah (#000000) - Ana renk, güven ve lüks ifade eder
• Beyaz (#FFFFFF) - Temizlik ve açıklık
• Gri (#F0F0F0, #808080) - İkincil yapı

Accent Colors:
• Altın (#D4AF37) - Vurgu ve özel işaretler
• Bej - Softwares ve arka plan tonları
• Yeşil (#22C55E) - Başarılı işlemler
• Kırmızı (#EF4444) - Uyarılar ve hatalar
""")

doc.add_heading("Tipografi (Typography)", level=3)
doc.add_paragraph("""
Font Ailesi: Sans-serif (Modern ve okunabilir)
- Başlıklar: Bold 24-32px
- Alt Başlıklar: Semibold 18-20px  
- Gövde Metni: Regular 14-16px
- Açıklamalar: Light 12px

Satır Yüksekliği: 1.5x - 1.6x
Harf Boşluğu: Standart
""")

doc.add_heading("Boşluk ve Alignement (Spacing)", level=3)
doc.add_paragraph("""
Base Unit: 4px (CSS rem/em kullanarak responsive)

Yaygın Boşluklar:
- Compact: 8px
- Normal: 16px  
- Generous: 24px
- Large: 32px

Alignment: Center, Left, Right - İçerik türüne göre
Padding: 16-32px
Margin: 8-24px
""")

doc.add_heading("Köşe Radii (Border Radius)", level=3)
doc.add_paragraph("""
- Sharp: 0px (Keskin köşeler - Minimalist)
- Small: 8px (Düğmeler ve inputs)
- Medium: 12px (Kartlar)
- Large: 16px (Konteynerler)
- Fully Rounded: 50% (Avatarlar, badge'ler)
""")

doc.add_heading("Gölgeler ve Derinlik (Elevation)", level=3)
doc.add_paragraph("""
Subtle Shadow: 0 1px 3px rgba(0,0,0,0.1)
Medium Shadow: 0 4px 6px rgba(0,0,0,0.1)
Large Shadow: 0 10px 15px rgba(0,0,0,0.1)

Hover State: +0 5px 8px rgba(0,0,0,0.15)
Active State: +0 2px 4px rgba(0,0,0,0.2)
""")

doc.add_heading("Animasyon ve Geçişler", level=3)
doc.add_paragraph("""
Transition Timings:
- Hızlı (Fast): 150-200ms - Hover effects, small changes
- Normal (Default): 300ms - Açma/kapama işlemleri
- Yavaş (Slow): 500ms+ - Kompleks animasyonlar, sayfa geçişleri

Easing Functions:
- ease-in-out: Standart geçişler
- ease-out: Elementlerin görünümü
- ease-in: Elementlerin gizlenmesi
- cubic-bezier: Custom animasyonlar (Framer Motion)
""")

doc.add_heading("İnteraktif Öğeler (Interactive Elements)", level=3)
doc.add_paragraph("""
Butonlar:
- Primary Button: Siyah arka plan, beyaz metin, 8px radius
- Secondary Button: Beyaz arka plan, siyah metin, border
- Size: 48px min height (mobile), padding: 12px 24px

Textbox ve Inputs:
- Border: 1px solid #E0E0E0
- Focus: Border rengi #000, shadow effect
- Hover: Arka plan: #F9F9F9
- Min height: 44px (mobile)

Links:
- Color: #000 (Siyah)
- Hover: Underline
- Active: Farklı renk

Navigation:
- Bottom Nav: Fixed position, 7 buton
- Each button: İkon + text label
- Active indicator: Farklı arka plan/renk
""")

doc.add_heading("Responsive Breakpoints", level=3)
doc.add_paragraph("""
Mobile: 320px - 767px
- Single column layout
- Larger touch targets (48x48px min)
- Bottom navigation

Tablet: 768px - 1023px
- Two column layout
- Balanced spacing
- Hybrid navigation

Desktop: 1024px+
- Multi-column layout
- Sidebar possible
- Advanced features
- Hover states
""")

doc.add_heading("Erişilebilirlik Standartları", level=3)
doc.add_paragraph("""
WCAG 2.1 Level AA Uyumluluğu:

Renk Kontrastı:
- Normal metin: 4.5:1
- Büyük metin: 3:1
- UI Bileşenleri: 3:1

Keyboard Navigation:
- Tab ile tüm elements'e erişim
- Enter ile aktivasyon
- Esc ile kapatma
- Arrow keys ile navigation

Screen Reader:
- Semantic HTML (button, link, heading, etc.)
- ARIA labels ve descriptions
- Alt text for images
- Form labels associated

Focus Management:
- Clear focus indicator (min 3px)
- Focus order logical
- Skip to main content link
""")

# Belgeyi kaydet
try:
    doc.save(updated_doc)
    print(f"✅ Başarılı!")
    print(f"📸 {added_count} görüntü eklendi")
    print(f"📄 Dosya: {updated_doc}")
    print(f"📊 Boyut: {os.path.getsize(updated_doc) / 1024:.2f} KB")
except Exception as e:
    print(f"❌ Hata: {e}")
