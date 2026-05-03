#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word belgesi güncelle - Proje görselleri ekle (v2)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_image_to_doc(doc, image_path, width=Inches(5.0)):
    """Resimleri dokümana ekle"""
    if os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=width)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception as e:
            print(f"Görüntü eklenirken hata: {image_path} - {e}")
            return False
    else:
        print(f"Görüntü bulunamadı: {image_path}")
        return False

# Güncellenen belgeyi aç
doc_path = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_UPDATED.docx"

try:
    doc = Document(doc_path)
    print(f"Belge açıldı: {doc_path}")
except Exception as e:
    print(f"Hata: {e}")
    exit(1)

# === GÖRSELLERI EKLE ===
doc.add_page_break()
doc.add_heading("EK: PROJE GÖRSELLERI (SCREENSHOTS)", level=1)

# 1. Fashion Assets
doc.add_heading("Proje Masaları", level=2)

assets = [
    (r"c:\Users\LENOVO\son bitirme\frontend\src\assets\chic_fashion_shoot_photo_1772833711579.png", 
     "Şık Moda Fotoğrafı - Landing Page Görseli"),
    (r"c:\Users\LENOVO\son bitirme\frontend\src\assets\fashion_avatar_render_1772833388818.png", 
     "Fashion Avatar Görseli - 3D Avatar Örneği"),
    (r"c:\Users\LENOVO\son bitirme\frontend\src\assets\quiet_luxury_wardrobe_bg_1772833697604.png", 
     "Quiet Luxury Wardrobe - Arka Plan"),
    (r"c:\Users\LENOVO\son bitirme\frontend\src\assets\premium_fashion_visual_1772833612094.png", 
     "Premium Moda Görseli - UI Tasarım Elemanı"),
    (r"c:\Users\LENOVO\son bitirme\frontend\src\assets\pinterest_fashion_couple_editorial_1772833854361.png", 
     "Pinterest Moda Editöryeli - Lookbook İlham"),
    (r"c:\Users\LENOVO\son bitirme\frontend\src\assets\luxury_wardrobe_bg_1772833374722.png", 
     "Lüks Wardrobe Arka Planı"),
]

for image_path, caption in assets:
    if os.path.exists(image_path):
        doc.add_paragraph()
        add_image_to_doc(doc, image_path, width=Inches(4.5))
        p = doc.add_paragraph(caption, style='Caption')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

# 2. Backend Upload Görselleri (Fashion Items)
doc.add_page_break()
doc.add_heading("Ürün Katalog Örnekleri", level=2)

uploads_path = r"c:\Users\LENOVO\son bitirme\backend\uploads"
if os.path.exists(uploads_path):
    # Thumbnailer görselleri
    thumb_count = 0
    for file in os.listdir(uploads_path):
        if file.startswith("thumb_") and file.endswith(".png"):
            file_path = os.path.join(uploads_path, file)
            if add_image_to_doc(doc, file_path, width=Inches(2.5)):
                thumb_count += 1
                if thumb_count % 3 == 0:
                    doc.add_paragraph()
    
    if thumb_count > 0:
        print(f"{thumb_count} ürün görseli eklendi")

# 3. Teknoloji Bilgileri
doc.add_page_break()
doc.add_heading("Teknik Özellikler", level=1)

# Database Schema
doc.add_heading("Veritabanı Modelleri", level=2)
doc.add_paragraph("""
Prisma ORM ile tanımlanmış temel modeller:
• User (Kullanıcı) - E-posta, şifre, profil bilgileri
• Wardrobe (Gardırop) - Kullanıcının kişisel koleksiyonu
• Garment (Kıyafet) - Kıyafet detayları (marka, renk, kategori)
• Outfit (Kombinasyon) - Kıyafet kombinleri
• Avatar (3D Avatar) - Kullanıcı 3D avatarı
• Recommendation (Öneriler) - AI tarafından oluşturulan öneriler
• Log (İşlem Kayıtları) - Sistem logları ve işlemleri
• RefreshToken - JWT refresh tokenları
""")

# Architecture
doc.add_heading("Proje Mimarisi", level=2)
doc.add_paragraph("""
Microservices yaklaşımıyla oluşturulan modüler yapı:

Frontend (React):
├── Components (Bileşenler)
├── Pages (Sayfalar)
├── Store (Zustand State Management)
├── Services (API çağrıları)
└── Assets (Görseller ve ikonlar)

Backend (NestJS):
├── Admin Module (Yönetici işlemleri)
├── Auth Module (Kimlik doğrulama)
├── Avatar Module (3D Avatar)
├── AI Module (Yapay Zeka özellikleri)
├── Wardrobe Module (Gardırop yönetimi)
├── Users Module (Kullanıcı yönetimi)
├── Social Module (Sosyal özellikler)
├── Outfits Module (Kombinasyonlar)
└── Infrastructure (Veritabanı, Storage, Queue)
""")

# Belgeyi kaydet
try:
    doc.save(doc_path)
    print(f"Belge başarılı bir şekilde güncellendi: {doc_path}")
    
    # Orijinal dosyayı yenile (backup + replace)
    import shutil
    original = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk.docx"
    backup = r"c:\Users\LENOVO\son bitirme\BitirmeOdevi_HasanOzturk_BACKUP.docx"
    
    if os.path.exists(original):
        shutil.copy2(original, backup)
        print(f"Yedek oluşturuldu: {backup}")
    
    shutil.copy2(doc_path, original)
    print(f"Orijinal dosya güncellendi: {original}")
    
except Exception as e:
    print(f"Hata: {e}")
